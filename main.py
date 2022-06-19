
import openpyxl
import docx
from openpyxl import load_workbook
import matplotlib.pyplot as plt
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QMainWindow

def start():
    app = QApplication(sys.argv)
    w = main_window()
    w.resize(700, 350)
    w.show()
    sys.exit(app.exec_())

def counting_columns(patch):
    wb = openpyxl.load_workbook(patch)
    sheets_list = wb.sheetnames
    sheet_active = wb[sheets_list[0]]
    row_max = sheet_active.max_row
    column_max = sheet_active.max_column

    print('В файле:', './test1.xlsx', '\n Cтолбцов:', row_max, '\n Колонок:', column_max)

class main_window(QMainWindow):

    def __init__(self):
        super(main_window, self).__init__()

        self.min = None
        self.max = None
        self.counter = 0
        self.minestr = ''
        self.cooler_currrent_1 = ''
        self.cooler_currrent_2 = ''
        self.device = ''
        self.dev_without_1 = ''
        self.dev_without_2 = ''

        self.textEdit = QTextEdit(self)
        self.textEdit.toPlainText()
        self.textEdit.setReadOnly(True)
        self.textEdit.hide()
        self.textEdit.resize(500, 250)

        self.btn1 = QPushButton("Выбор конфигурации", self)
        self.btn1.move(250, 150)
        self.btn1.resize(200, 50)

        self.lbl = QLabel(self)
        self.lbl.hide()
        self.qle = QLineEdit(self)
        self.qle.resize(500, 30)
        self.qle.hide()

        self.qle.move(0, 250)
        self.lbl.move(50, 50)

        self.btn2 = QPushButton("Загрузить файл", self)
        self.btn2.move(250, 150)
        self.btn2.resize(200, 50)

        self.btn3 = QPushButton('Start', self)
        self.btn3.move(150, 100)
        self.btn3.resize(50, 20)
        self.btn3.hide()

        self.btn4 = QPushButton('Тестовая конфигурация', self)
        self.btn4.move(500, 25)
        self.btn4.resize(200, 50)
        self.btn4.hide()

        self.btn5 = QPushButton('Конфигурация К53', self)
        self.btn5.move(500, 100)
        self.btn5.resize(200, 50)
        self.btn5.hide()

        self.btn1.clicked.connect(self.buttonClicked)
        self.btn2.clicked.connect(self.buttonClicked)
        self.btn3.clicked.connect(self.buttonClicked)
        self.btn4.clicked.connect(self.buttonClicked)
        self.btn5.clicked.connect(self.buttonClicked)

        self.btn3.clicked.connect(self.on_click)

        self.statusBar()

        self.setGeometry(300, 300, 290, 150)
        self.setWindowTitle('Программная система анализа производительности узлов персонального компьютера')
        self.show()
        self.btn2.hide()

    @pyqtSlot()
    def on_click(self):
        self.textEdit.append(str(self.qle.text()))
        self.qle.clear()

    def main(self, file):

        wb = load_workbook(file)
        sheet = wb['Лист1']

        for i in range(4, 752):

            if ((sheet.cell(row=i + 1, column=43).value - sheet.cell(row=i, column=43).value) < -400) \
                    and sheet.cell(row=i, column=44).value > 100:

                if sheet.cell(row=i + 1, column=4).value == 3:
                    data_string = f"В строке {i} резкое падание тока потребления с {sheet.cell(row=i, column=43).value} до {sheet.cell(row=i+1, column=43).value} ."
                    self.minestr = str(data_string)
                    self.textEdit.append(self.minestr)
                    data_string = 'Статус ошибки кулера:', sheet.cell(row=i+1, column=4).value
                    self.minestr = str(data_string)
                    self.textEdit.append(self.minestr)
                    self.cooler_currrent_1 = 'True'

                if sheet.cell(row=i + 1, column=4).value == 0:
                    data_string = f"В строке {i} резкое падание тока потребления с {sheet.cell(row=i, column=43).value} до {sheet.cell(row=i+1, column=43).value} ."
                    self.minestr = str(data_string)
                    self.textEdit.append(self.minestr)
                    data_string = 'Статус ошибки кулера:', sheet.cell(row=i+1, column=4).value
                    self.minestr = str(data_string)
                    self.textEdit.append(self.minestr)
                    self.cooler_currrent_2 = 'True'

                if sheet.cell(row=i, column=5).value == 0:
                    self.device = 'True'

                if sheet.cell(row=i, column=5).value == 0 and sheet.cell(row=i, column=44).value > 100:
                    print(sheet.cell(row=i, column=5).value)
                    self.dev_without_1 = 'True'
                if sheet.cell(row=i, column=5).value == 0 and sheet.cell(row=i, column=44).value < 100:
                    print(sheet.cell(row=i, column=5).value)
                    self.dev_without_2 = 'True'

        data_time = []  # Время
        for i in range(4, 753):
            data_time.append(sheet.cell(row=i, column=1).value)
        data_io = []  # Ток
        for i in range(4, 753):
            data_io.append(sheet.cell(row=i, column=43).value)
        data_fan1 = []  # Fan1
        for i in range(4, 753):
            data_fan1.append(sheet.cell(row=i, column=6).value)
        data_fan2 = []  # Fan2
        for i in range(4, 753):
            data_fan2.append(sheet.cell(row=i, column=7).value)
        data_cpu_vcore = []  # CPU Vcore
        for i in range(4, 753):
            data_cpu_vcore.append(sheet.cell(row=i, column=12).value)
        data_gpu_vcore = []  # GPU Vcore
        for i in range(4, 753):
            data_gpu_vcore.append(sheet.cell(row=i, column=13).value)

        self.report(self.cooler_currrent_1, self.cooler_currrent_2, self.device, self.dev_without_2, self.dev_without_2)

        data_io_min = self.min
        data_io_max = self.max
        fig = plt.figure("Графики")
        # subplot 1
        sp = plt.subplot(221)
        plt.title('Зависимость потребления тока от времени')
        plt.xlabel('Время, с', fontsize=14)
        plt.ylabel('Потребление тока, мА', fontsize=14)
        plt.grid(True)
        plt.plot(data_time, data_io)
        x = [0, 753]
        y = [data_io_min, data_io_min]
        plt.plot(x, y)
        y = [data_io_max, data_io_max]
        plt.plot(x, y)

        # subplot 2
        sp = plt.subplot(222)
        plt.title('Скорость вращения вентиляторов во времени')
        plt.xlabel('Время, с', fontsize=14)
        plt.ylabel('Скорость вращения вентиляторов', fontsize=14)
        plt.grid(True)
        plt.plot(data_time, data_fan1)  # Могут вести себя идентично и накладываться друг на друга
        plt.plot(data_time, data_fan2)
        x = [0, 753]
        y = [data_io_min, data_io_min]
        plt.plot(x, y)
        y = [data_io_max, data_io_max]
        plt.plot(x, y)

        # subplot 3
        sp = plt.subplot(223)
        plt.title('CPU')
        plt.xlabel('Время, с', fontsize=14)
        plt.ylabel('Потребление напряжения, мВ', fontsize=14)
        plt.grid(True)
        plt.plot(data_time, data_cpu_vcore)
        x = [0, 753]
        y = [data_io_min, data_io_min]
        plt.plot(x, y)
        y = [data_io_max, data_io_max]
        plt.plot(x, y)

        # subplot 4
        sp = plt.subplot(224)
        plt.title('GPU')
        plt.xlabel('Время, с', fontsize=14)
        plt.ylabel('Потребление напряжения, мВ', fontsize=14)
        plt.grid(True)
        plt.plot(data_time, data_gpu_vcore)
        x = [0, 753]
        y = [data_io_min, data_io_min]
        plt.plot(x, y)
        y = [data_io_max, data_io_max]
        plt.plot(x, y)


        plt.show()

    def report(self, cooler_currrent_1, cooler_currrent_2, device, dev1, dev2):

        if cooler_currrent_1 == True:
            report_doc = docx.Document()
            report_doc.add_paragraph("Обнаружено резкое падение тока потребления и ошибка кулера")
            report_doc.add_paragraph("Нужно почистить систему охлаждения")
            report_doc.save('E:/Report_doc.docx')

        if cooler_currrent_2 == True:
            report_doc = docx.Document()
            report_doc.add_paragraph("Обнаружено резкое падение тока потребления и кулер исправен")
            report_doc.add_paragraph("Почистите вентилятор от пыли или замените на более мощный")
            report_doc.save('E:/Report_doc.docx')

        if device == True:
            report_doc = docx.Document()
            report_doc.add_paragraph("Обнаружено отключенное устройство")
            report_doc.add_paragraph("Проверьте конфигурацию или возможно устройство не работает")
            report_doc.save('E:/Report_doc.docx')

        if dev1 == True:
            report_doc = docx.Document()
            report_doc.add_paragraph("Перегревается устройство без охлаждения")
            report_doc.add_paragraph("Добавьте охлаждение или настройте воздушные потоки")
            report_doc.save('E:/Report_doc.docx')

        if dev2 == True:
            report_doc = docx.Document()
            report_doc.add_paragraph("Отключение элемента без охлаждения")
            report_doc.add_paragraph("Скорее всего ус-во неисправно, т.к. перегрев не обнаружен")
            report_doc.save('E:/Report_doc.docx')



    def after(self, pr, st):
        self.hide()
        self.lbl.setText('pr + st')
        self.show()

    def buttonClicked(self):
        sender = self.sender()
        #self.statusBar().showMessage('Выберите файл')
        if sender.text() == 'Выбор конфигурации':
            self.btn1.hide()
            self.btn4.show()
            self.btn5.show()
            #self.btn2.show()

        if sender.text() == 'Конфигурация К53':
            self.msg = QMessageBox()
            self.msg.setWindowTitle("Оповещение")
            self.msg.setText("Конфигурация К53 еще в разработке")
            self.msg.setIcon(QMessageBox.Warning)
            self.msg.exec_()

        if sender.text() == 'Тестовая конфигурация':
            self.btn4.hide()
            self.btn5.hide()
            self.btn2.show()

        if sender.text() == 'Загрузить файл':
            self.file = QtWidgets.QFileDialog.getOpenFileName()[0]
            counting_columns(self.file)
            self.btn2.hide()
            self.lbl.show()
            self.qle.show()
            #self.btn3.show()
            self.textEdit.show()
            self.textEdit.append('Для выбранной конфигурации нужно выбрать минимальный и максимальный пороги.')
            self.textEdit.append('Укажите минимальный, а затем максимальный порог:')
            #self.main(wb_patch) # кнопку ок заменить на кнопку старт, запись значений с qle через enter фиксировать, а старт программы - start

        if sender.text() == 'Start':
            print('lel') # Скрыл кнопку старт, работает без нее, продумать, нужна ли она

    def keyPressEvent(self, e):

        if e.key() == Qt.Key_Escape:
            self.hide()

        if e.key() == Qt.Key_Return:
            number = int(self.qle.text())
            if self.max == None and self.min == None:
                self.min = int(self.qle.text())
                self.max = int(self.qle.text())
                self.counter += 1
                self.textEdit.append(self.qle.text())
                self.qle.clear()

            elif int(self.qle.text()) >= self.max:
                self.max = int(self.qle.text())
                self.counter += 1
                self.textEdit.append(self.qle.text())
                self.qle.clear()

            elif int(self.qle.text()) <= self.min:
                self.min = int(self.qle.text())
                self.counter += 1
                self.textEdit.append(self.qle.text())
                self.qle.clear()

            if self.counter == 2:
                self.main(self.file)

start()
